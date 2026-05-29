"""Resume template builder — generates filled DOCX from analysis data."""
from dataclasses import dataclass, field
from io import BytesIO
from typing import Optional

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement


@dataclass
class TemplateConfig:
    name: str
    font_body: str = "Calibri"
    font_heading: str = "Calibri Light"
    font_size_body: int = 11
    font_size_section: int = 13
    font_size_name: int = 22
    font_size_contact: int = 10
    color_primary: str = "2c3e50"
    color_accent: str = "2c3e50"
    color_text: str = "333333"
    header_align: str = "center"  # "center" | "left"
    section_style: str = "underline"  # "underline" | "border-bottom" | "line"
    show_section_icons: bool = False
    two_column: bool = False
    has_sidebar: bool = False
    page_margin_top: int = 60
    page_margin_bottom: int = 60
    page_margin_left: int = 60
    page_margin_right: int = 60


TEMPLATE_CONFIGS: dict[str, TemplateConfig] = {
    # Two-column sidebar + main, centered header, primary-colored section headings, timeline feel
    "azurill": TemplateConfig(
        name="Azurill",
        font_body="Calibri", font_heading="Calibri Light",
        color_primary="2563eb", color_accent="2563eb",
        header_align="center", section_style="underline",
        page_margin_top=50, page_margin_bottom=50,
    ),
    # Interleaved single-column, centered header, sections with border-top dividers
    "bronzor": TemplateConfig(
        name="Bronzor",
        font_body="Calibri", font_heading="Calibri",
        color_primary="0d9488", color_accent="0d9488",
        header_align="center", section_style="border-bottom",
        page_margin_top=55, page_margin_bottom=55,
    ),
    # Two-column with colored sidebar background, bold contrast, border-bottom headings
    "chikorita": TemplateConfig(
        name="Chikorita",
        font_body="Calibri", font_heading="Calibri Light",
        color_primary="166534", color_accent="166534",
        header_align="left", section_style="border-bottom",
        page_margin_top=50, page_margin_bottom=50,
    ),
    # Two-column with tinted sidebar, summary in special container, header in sidebar
    "ditgar": TemplateConfig(
        name="Ditgar",
        font_body="Calibri", font_heading="Calibri Light",
        color_primary="5b21b6", color_accent="5b21b6",
        header_align="left", section_style="underline",
        page_margin_top=50, page_margin_bottom=50,
    ),
    # Two-column with full-width header band, picture overlapping, no underline
    "ditto": TemplateConfig(
        name="Ditto",
        font_body="Segoe UI", font_heading="Segoe UI",
        color_primary="b45309", color_accent="d97706",
        header_align="center", section_style="underline",
        page_margin_top=50, page_margin_bottom=50,
    ),
    # Two-column with solid sidebar bg, featured summary, border-bottom headings
    "gengar": TemplateConfig(
        name="Gengar",
        font_body="Calibri", font_heading="Calibri Light",
        color_primary="581c87", color_accent="581c87",
        header_align="left", section_style="border-bottom",
        page_margin_top=50, page_margin_bottom=50,
    ),
    # Two-column with contact card in sidebar, centered header
    "glalie": TemplateConfig(
        name="Glalie",
        font_body="Calibri", font_heading="Calibri Light",
        color_primary="0369a1", color_accent="0284c7",
        header_align="center", section_style="border-bottom",
        page_margin_top=50, page_margin_bottom=50,
    ),
    # Single-column fully centered, minimalist, compact
    "kakuna": TemplateConfig(
        name="Kakuna",
        font_body="Calibri", font_heading="Calibri",
        color_primary="4b5563", color_accent="4b5563",
        header_align="center", section_style="underline",
        font_size_name=20, font_size_body=10, font_size_section=11,
        page_margin_top=40, page_margin_bottom=40, page_margin_left=70, page_margin_right=70,
    ),
    # Single-column card-based section boxes, side-by-side header
    "lapras": TemplateConfig(
        name="Lapras",
        font_body="Calibri", font_heading="Calibri Light",
        color_primary="1e40af", color_accent="3b82f6",
        header_align="left", section_style="underline",
        page_margin_top=50, page_margin_bottom=50,
    ),
    # Two-tone layered header bands, two-column
    "leafish": TemplateConfig(
        name="Leafish",
        font_body="Calibri", font_heading="Calibri Light",
        color_primary="065f46", color_accent="047857",
        header_align="center", section_style="underline",
        page_margin_top=50, page_margin_bottom=50,
    ),
    # Single-column, side-by-side header, uppercase section headings with border-bottom
    "meowth": TemplateConfig(
        name="Meowth",
        font_body="Calibri", font_heading="Calibri",
        color_primary="9a3412", color_accent="c2410c",
        header_align="left", section_style="border-bottom",
        page_margin_top=50, page_margin_bottom=50,
    ),
    # Single-column clean, side-by-side header with primary bottom border
    "onyx": TemplateConfig(
        name="Onyx",
        font_body="Calibri", font_heading="Calibri Light",
        color_primary="1e3a5f", color_accent="1e3a5f",
        header_align="left", section_style="underline",
        page_margin_top=55, page_margin_bottom=55,
    ),
    # Two-column with pill-shaped header in main column, picture in sidebar
    "pikachu": TemplateConfig(
        name="Pikachu",
        font_body="Calibri", font_heading="Calibri Light",
        color_primary="d97706", color_accent="d97706",
        header_align="center", section_style="border-bottom",
        page_margin_top=50, page_margin_bottom=50,
    ),
    # Single-column, dot-separated contacts, border-bottom headings
    "rhyhorn": TemplateConfig(
        name="Rhyhorn",
        font_body="Calibri", font_heading="Calibri",
        color_primary="475569", color_accent="64748b",
        header_align="left", section_style="border-bottom",
        page_margin_top=50, page_margin_bottom=50,
    ),
    # Single-column with top border, decorative rule under name, section dividers
    "scizor": TemplateConfig(
        name="Scizor",
        font_body="Calibri", font_heading="Calibri",
        color_primary="991b1b", color_accent="dc2626",
        header_align="left", section_style="underline",
        font_size_name=24,
        page_margin_top=55, page_margin_bottom=55,
    ),
}


@dataclass
class BuildResult:
    buffer: BytesIO
    estimated_pages: int
    overflow_warning: bool
    template_page_limit: Optional[int]
    extra_sections_appended: list[str]


def _hex_to_rgb(hex_color: str) -> RGBColor:
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _set_cell_shading(cell, color: str):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _add_border_bottom(paragraph, color_rgb: RGBColor, sz: int = 6):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), str(color_rgb))
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_paragraph_spacing(paragraph, before: int = 0, after: int = 4, line: int | None = None):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line:
        pf.line_spacing = Pt(line)


def _add_run(paragraph, text: str, font_name: str, font_size: int,
             bold: bool = False, color: RGBColor | None = None,
             italic: bool = False):
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    run.font.italic = italic
    return run


def _apply_document_defaults(doc: Document, config: TemplateConfig):
    """Set default font, margins, and spacing for the document."""
    section = doc.sections[0]
    section.top_margin = Pt(config.page_margin_top)
    section.bottom_margin = Pt(config.page_margin_bottom)
    section.left_margin = Pt(config.page_margin_left)
    section.right_margin = Pt(config.page_margin_right)

    style = doc.styles["Normal"]
    style.font.name = config.font_body
    style.font.size = Pt(config.font_size_body)
    style.font.color.rgb = _hex_to_rgb(config.color_text)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.space_before = Pt(0)


def _render_header(doc: Document, resume_data: dict, config: TemplateConfig):
    """Render name, contact info, summary."""
    name = resume_data.get("name") or "Your Name"
    email = resume_data.get("email")
    phone = resume_data.get("phone")
    linkedin = resume_data.get("linkedin")
    github = resume_data.get("github")
    location = resume_data.get("location")
    summary = resume_data.get("summary")

    align = WD_ALIGN_PARAGRAPH.CENTER if config.header_align == "center" else WD_ALIGN_PARAGRAPH.LEFT

    # Name
    p = doc.add_paragraph()
    p.alignment = align
    _set_paragraph_spacing(p, before=0, after=2)
    _add_run(p, name, config.font_heading, config.font_size_name, bold=True,
             color=_hex_to_rgb(config.color_primary))

    # Contact line
    contact_parts = [email, phone, location, linkedin, github]
    contact_text = "  |  ".join(p for p in contact_parts if p)
    if contact_text:
        p2 = doc.add_paragraph()
        p2.alignment = align
        _set_paragraph_spacing(p2, before=0, after=4)
        _add_run(p2, contact_text, config.font_body, config.font_size_contact,
                 color=_hex_to_rgb(config.color_accent))

    # Summary
    if summary:
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_paragraph_spacing(p3, before=4, after=6)
        _add_run(p3, summary, config.font_body, config.font_size_body,
                 italic=True)


def _render_section_heading(doc: Document, title: str, config: TemplateConfig):
    """Render a section heading with the configured style."""
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, before=8, after=4)
    primary_rgb = _hex_to_rgb(config.color_accent)

    if config.section_style == "underline":
        run = _add_run(p, title.upper(), config.font_heading, config.font_size_section,
                       bold=True, color=primary_rgb)
        run.font.underline = True
    elif config.section_style == "border-bottom":
        _add_run(p, title.upper(), config.font_heading, config.font_size_section,
                 bold=True, color=primary_rgb)
        _add_border_bottom(p, primary_rgb, sz=8)
    elif config.section_style == "line":
        _add_run(p, title.upper(), config.font_heading, config.font_size_section,
                 bold=True, color=primary_rgb)
        line_p = doc.add_paragraph()
        _set_paragraph_spacing(line_p, before=0, after=4)
        run = line_p.add_run("─" * 60)
        run.font.color.rgb = primary_rgb
        run.font.size = Pt(6)


def _render_experience_section(doc: Document, experiences: list, rephrase_map: dict,
                                config: TemplateConfig):
    """Render work experience entries with rephrased bullets."""
    if not experiences:
        return
    _render_section_heading(doc, "Experience", config)

    for exp_idx, exp in enumerate(experiences):
        company = exp.get("company") or ""
        title = exp.get("title") or ""
        start = exp.get("start_date") or ""
        end = exp.get("end_date") or ""
        descriptions = exp.get("descriptions") or []

        # Company / Title / Dates line
        p = doc.add_paragraph()
        _set_paragraph_spacing(p, before=2, after=1)
        if company and title:
            _add_run(p, f"{company} — {title}", config.font_body, config.font_size_body,
                     bold=True, color=_hex_to_rgb(config.color_primary))
        elif company:
            _add_run(p, company, config.font_body, config.font_size_body,
                     bold=True)
        elif title:
            _add_run(p, title, config.font_body, config.font_size_body,
                     bold=True)

        if start or end:
            p2 = doc.add_paragraph()
            _set_paragraph_spacing(p2, before=0, after=2)
            date_str = f"{start} — {end}" if start and end else (start or end)
            _add_run(p2, date_str, config.font_body, config.font_size_body - 1,
                     color=_hex_to_rgb("666666"))

        # Bullet points
        for b_idx, bullet in enumerate(descriptions):
            key = f"experience__{exp_idx}__{b_idx}"
            content = rephrase_map.get(key, bullet)
            if content:
                pb = doc.add_paragraph(style="List Bullet")
                _set_paragraph_spacing(pb, before=0, after=1)
                _add_run(pb, content, config.font_body, config.font_size_body)

        # Small gap between entries
        if exp_idx < len(experiences) - 1:
            spacer = doc.add_paragraph()
            _set_paragraph_spacing(spacer, before=0, after=2)


def _render_education_section(doc: Document, education: list, config: TemplateConfig):
    if not education:
        return
    _render_section_heading(doc, "Education", config)

    for edu in education:
        institution = edu.get("institution") or edu.get("name") or ""
        degree = edu.get("degree") or ""
        score_val = edu.get("score") or ""
        start = edu.get("start_date") or ""
        end = edu.get("end_date") or ""
        location = edu.get("location") or ""

        p = doc.add_paragraph()
        _set_paragraph_spacing(p, before=2, after=1)
        parts = [p for p in [institution, degree] if p]
        _add_run(p, " — ".join(parts), config.font_body, config.font_size_body,
                 bold=True, color=_hex_to_rgb(config.color_primary))

        details = [p for p in [start, end] if p]
        detail_str = f"{' — '.join(details)}" if details else ""
        if score_val:
            detail_str = f"{detail_str} | GPA: {score_val}" if detail_str else f"GPA: {score_val}"
        if detail_str:
            p2 = doc.add_paragraph()
            _set_paragraph_spacing(p2, before=0, after=2)
            _add_run(p2, detail_str, config.font_body, config.font_size_body - 1,
                     color=_hex_to_rgb("666666"))


def _render_skills_section(doc: Document, skills: list, config: TemplateConfig):
    if not skills:
        return
    _render_section_heading(doc, "Skills", config)
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, before=2, after=2)
    # Display as comma-separated inline
    _add_run(p, ", ".join(skills), config.font_body, config.font_size_body)


def _render_projects_section(doc: Document, projects: list, rephrase_map: dict,
                              config: TemplateConfig):
    if not projects:
        return
    _render_section_heading(doc, "Projects", config)

    for proj_idx, proj in enumerate(projects):
        name = proj.get("name") or ""
        descriptions = proj.get("descriptions") or []
        link = proj.get("link") or ""

        p = doc.add_paragraph()
        _set_paragraph_spacing(p, before=2, after=1)
        _add_run(p, name, config.font_body, config.font_size_body,
                 bold=True, color=_hex_to_rgb(config.color_primary))

        if link:
            p2 = doc.add_paragraph()
            _set_paragraph_spacing(p2, before=0, after=1)
            _add_run(p2, link, config.font_body, config.font_size_body - 1,
                     color=_hex_to_rgb(config.color_accent))

        for b_idx, bullet in enumerate(descriptions):
            key = f"projects__{proj_idx}__{b_idx}"
            content = rephrase_map.get(key, bullet)
            if content:
                pb = doc.add_paragraph(style="List Bullet")
                _set_paragraph_spacing(pb, before=0, after=1)
                _add_run(pb, content, config.font_body, config.font_size_body)

        if proj_idx < len(projects) - 1:
            spacer = doc.add_paragraph()
            _set_paragraph_spacing(spacer, before=0, after=2)


def _render_certifications_section(doc: Document, certifications: list, config: TemplateConfig):
    if not certifications:
        return
    _render_section_heading(doc, "Certifications", config)
    for cert in certifications:
        name = cert.get("name") or ""
        issuer = cert.get("issuer") or ""
        date_str = cert.get("date") or ""
        parts = [p for p in [name, issuer] if p]
        text = f"{' — '.join(parts)}" + (f" ({date_str})" if date_str else "")
        if text:
            p = doc.add_paragraph(style="List Bullet")
            _set_paragraph_spacing(p, before=0, after=1)
            _add_run(p, text, config.font_body, config.font_size_body)


def _render_simple_section(doc: Document, title: str, items: list,
                           config: TemplateConfig):
    """Generic section renderer for simple list-based sections."""
    if not items:
        return
    _render_section_heading(doc, title, config)
    for item in items:
        if isinstance(item, str):
            p = doc.add_paragraph(style="List Bullet")
            _set_paragraph_spacing(p, before=0, after=1)
            _add_run(p, item, config.font_body, config.font_size_body)
        elif isinstance(item, dict):
            title_text = item.get("title") or item.get("name", "")
            descriptions = item.get("descriptions") or []
            if title_text:
                p = doc.add_paragraph()
                _set_paragraph_spacing(p, before=2, after=1)
                _add_run(p, title_text, config.font_body, config.font_size_body,
                         bold=True, color=_hex_to_rgb(config.color_primary))
            for desc in descriptions:
                pb = doc.add_paragraph(style="List Bullet")
                _set_paragraph_spacing(pb, before=0, after=1)
                _add_run(pb, desc, config.font_body, config.font_size_body)


def _append_extra_sections(doc: Document, resume_data: dict,
                           config: TemplateConfig,
                           supported_sections: set[str]) -> list[str]:
    """Append sections present in resume data but not supported by template."""
    EXTRA_MAP = {
        "achievements": "Achievements",
        "hobbies": "Hobbies & Interests",
        "extra_curricular": "Extracurricular",
    }

    appended: list[str] = []

    for field_key, heading in EXTRA_MAP.items():
        section_name = heading
        if section_name in supported_sections:
            continue
        items = resume_data.get(field_key, [])
        if items:
            _render_simple_section(doc, section_name, items, config)
            appended.append(section_name)

    return appended


def _estimate_page_count(doc: Document) -> int:
    """Rough page count estimation based on content density."""
    total_paragraphs = len(doc.paragraphs)
    total_chars = sum(len(p.text) for p in doc.paragraphs)

    total_tables = len(doc.tables)
    table_rows = sum(len(t.rows) for t in doc.tables)

    # Rough heuristic: ~45 lines per A4 page at 11pt with default margins
    line_estimate = total_paragraphs + table_rows + (total_chars / 80)
    return max(1, round(line_estimate / 40))


def build_resume(template_id: str, resume_data: dict,
                 rephrase_map: dict[str, str] | None = None,
                 supported_sections: list[str] | None = None,
                 page_limit: int | None = None) -> BuildResult:
    """Generate a filled DOCX from analysis data using the given template."""
    config = TEMPLATE_CONFIGS.get(template_id)
    if not config:
        available = list(TEMPLATE_CONFIGS.keys())
        raise ValueError(f"Unknown template '{template_id}'. Available: {available}")

    rephrase_map = rephrase_map or {}
    supported = set(supported_sections or [])

    doc = Document()
    _apply_document_defaults(doc, config)

    # --- Render sections ---
    _render_header(doc, resume_data, config)

    _render_experience_section(doc, resume_data.get("experience", []), rephrase_map, config)

    _render_education_section(doc, resume_data.get("education", []), config)

    _render_skills_section(doc, resume_data.get("skills", []), config)

    _render_projects_section(doc, resume_data.get("projects", []), rephrase_map, config)

    _render_certifications_section(doc, resume_data.get("certifications", []), config)

    # Sections in resume that the template doesn't natively handle
    extra = _append_extra_sections(doc, resume_data, config, supported)

    estimated_pages = _estimate_page_count(doc)
    overflow = bool(page_limit) and estimated_pages > (page_limit or 999)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    return BuildResult(
        buffer=buf,
        estimated_pages=estimated_pages,
        overflow_warning=overflow,
        template_page_limit=page_limit,
        extra_sections_appended=extra,
    )
