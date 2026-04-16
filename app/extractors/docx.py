from docx import Document


def extract_docx(file_path: str) -> str:
    doc = Document(file_path)
    lines = []

    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            lines.append(" | ".join(cells))

    return "\n\n".join(lines)
