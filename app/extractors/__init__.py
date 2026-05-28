import os

import pymupdf
import pymupdf4llm
from docx import Document

__all__ = ["extract", "ResumeTooLongError"]

MAX_PAGES = 4


class ResumeTooLongError(Exception):
    def __init__(self, pages: int, max_pages: int) -> None:
        self.pages = pages
        self.max_pages = max_pages
        super().__init__(f"Resume exceeds {max_pages} pages ({pages} pages)")


def _extract_pdf(file_path: str) -> str:
    doc = pymupdf.open(file_path)
    page_count = len(doc)
    # Get raw text with sorting for better layout preservation in fallback
    raw_text = "\n\n".join(page.get_text(sort=True).strip() for page in doc)
    doc.close()

    if page_count > MAX_PAGES:
        raise ResumeTooLongError(page_count, MAX_PAGES)

    try:
        md_text = pymupdf4llm.to_markdown(file_path)
    except Exception:
        return raw_text

    # Check for missing critical info (like email) in markdown that exists in raw text
    has_email_raw = "@" in raw_text
    has_email_md = "@" in md_text
    
    # If markdown is missing an email that was in the raw text, or is significantly shorter
    # (indicating it might have skipped sections/tables), fall back to raw text.
    # LLMs handle raw text well, so it's safer than missing data.
    if (has_email_raw and not has_email_md) or (len(md_text) < len(raw_text) * 0.85):
        return raw_text

    return md_text


def _extract_docx(file_path: str) -> str:
    doc = Document(file_path)
    lines: list[str] = []

    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            lines.append(" | ".join(cells))

    return "\n\n".join(lines)


def extract(file_path: str) -> str:
    ext = file_path.lower().rsplit(".", 1)[-1]

    if ext == "pdf":
        return _extract_pdf(file_path)

    if ext == "docx":
        return _extract_docx(file_path)

    raise ValueError(f"Unsupported format: .{ext}")
