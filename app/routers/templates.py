from fastapi import APIRouter, HTTPException, Depends
from fastapi.responses import StreamingResponse
from io import BytesIO
from typing import Optional
from pydantic import BaseModel

from ..template_service import list_templates, get_template
from ..template_builder import build_resume
from .auth import get_current_user

router = APIRouter(prefix="/api/templates", tags=["templates"])


class BuildRequest(BaseModel):
    analysis_id: str


@router.get("")
async def get_templates():
    return list_templates()


@router.get("/{template_id}/preview")
async def preview_template(template_id: str):
    template = get_template(template_id)
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")
    return {"template_id": template_id, "preview": "Dummy preview data"}


@router.get("/{template_id}/download")
async def download_template(template_id: str):
    """Legacy download endpoint — returns unfilled template."""
    template = get_template(template_id)
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")

    from docx import Document
    doc = Document()
    doc.add_heading(template["name"], level=0)
    doc.add_paragraph("This resume was built with RateMyResume.")
    doc.add_heading("Experience", level=1)
    doc.add_paragraph("Your experience entries will appear here.")
    doc.add_heading("Education", level=1)
    doc.add_paragraph("Your education details will appear here.")
    doc.add_heading("Skills", level=1)
    doc.add_paragraph("Your skills will appear here.")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={template_id}.docx"},
    )


@router.post("/{template_id}/build")
async def build_resume_from_analysis(
    template_id: str,
    body: BuildRequest,
    current_user: dict = Depends(get_current_user),
):
    """Build a filled resume from an analysis result."""
    template = get_template(template_id)
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")

    # Load analysis from history
    from app.repositories.analysis import SupabaseAnalysisRepository
    from app.db import service_supabase

    repo = SupabaseAnalysisRepository(service_supabase)
    record = await repo.get_by_id(body.analysis_id)
    if not record:
        raise HTTPException(status_code=404, detail="Analysis not found")

    if record.get("user_id") != current_user["id"]:
        raise HTTPException(status_code=403, detail="Access denied")

    result_json = record.get("result_json")
    if not result_json:
        raise HTTPException(status_code=400, detail="Analysis has no result data")

    resume_data = result_json.get("resume_data") if isinstance(result_json, dict) else None
    if not resume_data:
        raise HTTPException(status_code=400, detail="Analysis has no resume data (re-analyze with current version)")

    rephrase_map = result_json.get("rephrase_map", {}) if isinstance(result_json, dict) else {}

    try:
        result = build_resume(
            template_id=template_id,
            resume_data=resume_data,
            rephrase_map=rephrase_map,
            supported_sections=template.get("supported_sections"),
            page_limit=template.get("page_limit"),
        )
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))

    filename = f"{template_id}_{body.analysis_id[:8]}.docx"
    headers = {
        "Content-Disposition": f"attachment; filename={filename}",
        "X-Estimated-Pages": str(result.estimated_pages),
        "X-Overflow-Warning": "true" if result.overflow_warning else "false",
        "X-Template-Limit": str(result.template_page_limit or ""),
        "X-Extra-Sections": ",".join(result.extra_sections_appended),
    }

    return StreamingResponse(
        result.buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )
