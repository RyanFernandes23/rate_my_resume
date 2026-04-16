import pytest
from docx import Document
import pymupdf

from app.extractors import extract


@pytest.fixture
def sample_pdf(tmp_path):
    doc = pymupdf.open()
    page = doc.new_page()
    page.insert_text(
        (50, 50),
        "John Doe\nEmail: john@example.com\n\nExperience\n- Developer at Tech Corp",
    )
    pdf_path = tmp_path / "resume.pdf"
    doc.save(str(pdf_path))
    doc.close()
    return str(pdf_path)


@pytest.fixture
def sample_docx(tmp_path):
    doc = Document()
    doc.add_heading("John Doe", level=1)
    doc.add_paragraph("Email: john@example.com")
    doc.add_heading("Experience", level=2)
    doc.add_paragraph("- Developer at Tech Corp")

    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Skill"
    table.cell(0, 1).text = "Level"
    table.cell(1, 0).text = "Python"
    table.cell(1, 1).text = "Expert"

    docx_path = tmp_path / "resume.docx"
    doc.save(str(docx_path))
    return str(docx_path)


class TestExtractPDF:
    def test_extract_pdf_returns_markdown(self, sample_pdf):
        result = extract(sample_pdf)
        assert isinstance(result, str)
        assert "John Doe" in result
        assert "john@example.com" in result


class TestExtractDOCX:
    def test_extract_docx_returns_markdown(self, sample_docx):
        result = extract(sample_docx)
        assert isinstance(result, str)
        assert "John Doe" in result
        assert "john@example.com" in result


class TestExtractUnsupported:
    def test_unsupported_format_raises_error(self, tmp_path):
        txt_path = tmp_path / "resume.txt"
        txt_path.write_text("test")

        with pytest.raises(ValueError) as exc_info:
            extract(str(txt_path))

        assert "Unsupported format" in str(exc_info.value)


class TestExtractFactory:
    def test_pdf_extension(self, tmp_path):
        pdf_path = tmp_path / "test.pdf"
        pdf_path.write_bytes(b"%PDF-1.4 test")

        with pytest.raises(Exception):
            extract(str(pdf_path))

    def test_docx_extension(self, tmp_path):
        docx_path = tmp_path / "test.docx"
        docx_path.write_bytes(b"PK\x03\x04")

        with pytest.raises(Exception):
            extract(str(docx_path))
